"""DOCX 抽取与 MIME 路由测试。"""
from __future__ import annotations

import io

import pytest
from docx import Document
from docx.shared import Inches

from app.core.ingest.errors import EMPTY_CONTENT, IngestError
from app.core.ingest.extractors.docx_extractor import DOCX_MIME, DocxExtractor
from app.core.ingest.extractors.file import FileExtractor, SUPPORTED_MIME_TYPES
from app.core.ingest.extractors.omml_to_text import omml_xml_to_text
from app.core.ocr.protocol import OcrRecognizeResult, OcrRegionDraft
from app.settings import get_settings


class FakeOcr:
    name = "fake"
    enabled = True
    provider_name = "fake"

    def __init__(self, text: str = "WORD_OCR") -> None:
        self.text = text

    async def recognize(self, image_bytes, *, lang=None, fail_open=None):
        return OcrRecognizeResult(text=self.text, confidence=0.8, empty=not self.text)

    async def recognize_region(self, *, page_no, bbox, crop_png, lang=None):
        return OcrRegionDraft(
            page_no=page_no,
            bbox=bbox,
            crop_png=crop_png,
            ocr_text=self.text,
            status="pending",
            provider="fake",
            confidence=0.8,
        )


def _tiny_png() -> bytes:
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (64, 48), color=(240, 240, 240)).save(buf, format="PNG")
    return buf.getvalue()


def _make_docx_bytes(*, text: str = "Hello Word", with_image: bool = False) -> bytes:
    doc = Document()
    doc.add_paragraph("Intro before image")
    if with_image:
        stream = io.BytesIO(_tiny_png())
        doc.add_picture(stream, width=Inches(1.0))
    doc.add_paragraph(text)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def test_docx_in_supported_mimes():
    assert DOCX_MIME in SUPPORTED_MIME_TYPES


def test_file_router_docx_magic():
    data = _make_docx_bytes()
    doc = FileExtractor().extract(data, original_filename="a.docx")
    assert doc.mime == DOCX_MIME
    assert "Hello Word" in doc.text
    assert "A | B" in doc.text


@pytest.mark.asyncio
async def test_docx_plain_extract():
    data = _make_docx_bytes(text="正文内容")
    doc = await DocxExtractor().extract_async(data, original_filename="t.docx")
    assert "正文内容" in doc.text
    assert "Intro before image" in doc.text


@pytest.mark.asyncio
async def test_docx_image_ocr_neighbor_prefix(monkeypatch):
    monkeypatch.setenv("RAG_OCR_PROVIDER", "api")
    get_settings.cache_clear()
    data = _make_docx_bytes(text="After image text", with_image=True)
    doc = await DocxExtractor(ocr_engine=FakeOcr("pic-text")).extract_async(
        data, original_filename="img.docx"
    )
    assert "[OCR image 1]" in doc.text
    assert "pic-text" in doc.text
    assert "Intro before image" in doc.text
    assert "After image text" in doc.text
    # 邻近：OCR 应在 Intro 之后、After 之前
    assert doc.text.index("Intro before image") < doc.text.index("[OCR image 1]")
    assert doc.text.index("[OCR image 1]") < doc.text.index("After image text")
    assert doc.metadata["ocr_region_drafts"]


@pytest.mark.asyncio
async def test_docx_without_ocr_keeps_text(monkeypatch):
    monkeypatch.setenv("RAG_OCR_PROVIDER", "none")
    get_settings.cache_clear()
    data = _make_docx_bytes(with_image=True)
    doc = await DocxExtractor().extract_async(data, original_filename="img.docx")
    assert "Hello Word" in doc.text
    assert "[OCR image" not in doc.text
    assert doc.metadata["ocr_region_drafts"] == []


def test_omml_fraction():
    xml = (
        b'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        b'<m:f><m:num><m:r><m:t>a</m:t></m:r></m:num>'
        b'<m:den><m:r><m:t>b</m:t></m:r></m:den></m:f></m:oMath>'
    )
    assert "(a)/(b)" in omml_xml_to_text(xml)


@pytest.mark.asyncio
async def test_empty_docx_fails():
    doc = Document()
    buf = io.BytesIO()
    doc.save(buf)
    with pytest.raises(IngestError) as ei:
        await DocxExtractor().extract_async(buf.getvalue(), original_filename="e.docx")
    assert ei.value.code == EMPTY_CONTENT
