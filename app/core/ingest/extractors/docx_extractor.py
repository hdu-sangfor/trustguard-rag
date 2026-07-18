"""Word (.docx) 抽取：正文/表格 + 嵌入原图 OCR + OMML/图表文本化。"""
from __future__ import annotations

import hashlib
import io
import logging
import zipfile
from typing import Any, Iterator

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.core.ingest.errors import (
    CORRUPT_FILE,
    DOCX_ENCRYPTED,
    EMPTY_CONTENT,
    FILE_TOO_LARGE,
    OCR_UNAVAILABLE,
    IngestError,
)
from app.core.ingest.extractors._async_utils import run_sync
from app.core.ingest.extractors.chart_text import chart_xml_to_text
from app.core.ingest.extractors.ocr_merge import format_ocr_span
from app.core.ingest.extractors.omml_to_text import omml_element_to_text
from app.core.ingest.models import ExtractedDocument
from app.core.ocr import get_ocr_engine
from app.core.ocr.errors import OcrError
from app.core.ocr.protocol import OcrRegionDraft
from app.settings import get_settings

logger = logging.getLogger(__name__)

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_A_BLIP = qn("a:blip")
_R_EMBED = qn("r:embed")
_M_OMATH = "{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath"


def _iter_block_items(document: DocumentObject) -> Iterator[Paragraph | Table]:
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _table_to_text(table: Table) -> str:
    rows: list[str] = []
    for row in table.rows:
        cells = [" ".join(c.text.split()) for c in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)


def _blip_rids(paragraph: Paragraph) -> list[str]:
    rids: list[str] = []
    for blip in paragraph._element.findall(".//" + _A_BLIP):
        rid = blip.get(_R_EMBED)
        if rid:
            rids.append(rid)
    return rids


def _omath_texts(paragraph: Paragraph) -> list[str]:
    out: list[str] = []
    for node in paragraph._element.findall(".//" + _M_OMATH):
        text = omml_element_to_text(node).strip()
        out.append(text if text else "[公式]")
    return out


def _is_docx_zip(data: bytes) -> bool:
    if not data.startswith(b"PK"):
        return False
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            names = set(zf.namelist())
            if "[Content_Types].xml" not in names:
                return False
            ct = zf.read("[Content_Types].xml").decode("utf-8", errors="ignore")
            return "wordprocessingml" in ct or "word/" in "".join(list(names)[:20])
    except zipfile.BadZipFile:
        return False


def _image_to_png(raw: bytes, content_type: str | None = None) -> bytes:
    """尽量转为 PNG 供 OCR；失败则原样返回。"""
    try:
        from PIL import Image

        img = Image.open(io.BytesIO(raw))
        if img.mode not in {"RGB", "L"}:
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    except Exception:  # noqa: BLE001
        return raw


def _extract_chart_blocks(data: bytes) -> list[str]:
    blocks: list[str] = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            for name in sorted(zf.namelist()):
                if not name.startswith("word/charts/chart") or not name.endswith(".xml"):
                    continue
                text = chart_xml_to_text(zf.read(name))
                if text:
                    blocks.append(text)
    except zipfile.BadZipFile:
        pass
    return blocks


def _header_footer_text(document: DocumentObject) -> tuple[str, str]:
    headers: list[str] = []
    footers: list[str] = []
    for section in document.sections:
        try:
            ht = "\n".join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
            if ht:
                headers.append(ht)
        except Exception:  # noqa: BLE001
            pass
        try:
            ft = "\n".join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
            if ft:
                footers.append(ft)
        except Exception:  # noqa: BLE001
            pass
    h = "\n".join(headers).strip()
    f = "\n".join(footers).strip()
    return h, f


class DocxExtractor:
    def __init__(self, ocr_engine=None) -> None:
        self._ocr = ocr_engine

    def extract(
        self,
        data: bytes,
        *,
        original_filename: str = "document.docx",
    ) -> ExtractedDocument:
        return run_sync(self.extract_async(data, original_filename=original_filename))

    async def extract_async(
        self,
        data: bytes,
        *,
        original_filename: str = "document.docx",
    ) -> ExtractedDocument:
        settings = get_settings()
        if len(data) > settings.ingest_max_file_bytes:
            raise IngestError(FILE_TOO_LARGE, "File exceeds max size")
        if not _is_docx_zip(data):
            raise IngestError(CORRUPT_FILE, "Not a valid DOCX file")

        # 粗检加密：EncryptedPackage
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                if "EncryptedPackage" in zf.namelist():
                    raise IngestError(DOCX_ENCRYPTED, "Password-protected DOCX not supported")
        except zipfile.BadZipFile as e:
            raise IngestError(CORRUPT_FILE, f"Cannot open DOCX: {e}") from e

        try:
            document = Document(io.BytesIO(data))
        except Exception as e:
            msg = str(e).lower()
            if "password" in msg or "encrypt" in msg:
                raise IngestError(DOCX_ENCRYPTED, "Password-protected DOCX not supported") from e
            raise IngestError(CORRUPT_FILE, f"Cannot open DOCX: {e}") from e

        ocr_engine = self._ocr or get_ocr_engine()
        region_drafts: list[OcrRegionDraft] = []
        body_parts: list[str] = []
        base_parts: list[str] = []
        image_no = 0
        max_doc_regions = settings.ocr_max_regions_per_document
        min_side = settings.ocr_min_image_side_px
        max_crop_bytes = settings.ocr_max_crop_bytes

        header, footer = _header_footer_text(document)
        if header:
            block = f"[页眉]\n{header}"
            body_parts.append(block)
            base_parts.append(block)

        for block in _iter_block_items(document):
            if isinstance(block, Table):
                t = _table_to_text(block)
                if t.strip():
                    body_parts.append(t)
                    base_parts.append(t)
                continue

            para = block
            # 段落正文
            para_text = (para.text or "").strip()
            if para_text:
                body_parts.append(para_text)
                base_parts.append(para_text)

            # 公式
            for formula in _omath_texts(para):
                formula_line = f"[公式] {formula}" if not formula.startswith("[") else formula
                body_parts.append(formula_line)
                base_parts.append(formula_line)

            # 嵌入图：原图 OCR，邻近插入
            for rid in _blip_rids(para):
                if len(region_drafts) >= max_doc_regions:
                    logger.warning("OCR region cap reached for document (%s)", max_doc_regions)
                    break
                try:
                    part = document.part.related_parts[rid]
                    raw = part.blob
                    content_type = getattr(part, "content_type", None)
                except Exception as e:  # noqa: BLE001
                    logger.warning("failed to load docx image %s: %s", rid, e)
                    continue
                if not raw:
                    continue

                image_no += 1
                png = _image_to_png(raw, content_type)
                # 尺寸粗检：无法解码时仍 OCR
                try:
                    from PIL import Image

                    im = Image.open(io.BytesIO(png))
                    w, h = im.size
                    if w < min_side or h < min_side:
                        continue
                    if w * h > settings.ocr_max_crop_pixels:
                        logger.warning(
                            "skip docx image exceeding max pixels size=%sx%s", w, h
                        )
                        continue
                    bbox = [0.0, 0.0, float(w), float(h)]
                except Exception:  # noqa: BLE001
                    bbox = [0.0, 0.0, 0.0, 0.0]
                if len(png) > max_crop_bytes:
                    logger.warning("skip docx image exceeding max bytes size=%s", len(png))
                    continue

                if not ocr_engine.enabled:
                    # OCR 关闭：跳过识别，保留正文
                    continue

                try:
                    draft = await ocr_engine.recognize_region(
                        page_no=None,
                        bbox=bbox,
                        crop_png=png,
                    )
                except OcrError as e:
                    if settings.ocr_fail_open:
                        draft = OcrRegionDraft(
                            page_no=None,
                            bbox=bbox,
                            crop_png=png,
                            ocr_text="",
                            status="failed",
                            provider=ocr_engine.provider_name,
                            error_message=str(e)[:200],
                            metadata={
                                "source": "docx",
                                "rel_id": rid,
                                "image_no": image_no,
                                "sequence": len(region_drafts),
                            },
                        )
                    else:
                        raise IngestError(OCR_UNAVAILABLE, str(e)) from e
                draft.metadata = {
                    **(draft.metadata or {}),
                    "source": "docx",
                    "rel_id": rid,
                    "image_no": image_no,
                    "content_type": content_type,
                    "sequence": len(region_drafts),
                }
                region_drafts.append(draft)
                span = format_ocr_span(image_no, draft.ocr_text)
                if span:
                    body_parts.append(span)

        for chart_block in _extract_chart_blocks(data):
            body_parts.append(chart_block)
            base_parts.append(chart_block)

        if footer:
            block = f"[页脚]\n{footer}"
            body_parts.append(block)
            base_parts.append(block)

        full_text = "\n\n".join(p for p in body_parts if p and p.strip()).strip()
        base_text = "\n\n".join(p for p in base_parts if p and p.strip()).strip()
        if not full_text and not region_drafts:
            raise IngestError(EMPTY_CONTENT, "DOCX has no extractable content")

        content_hash = hashlib.sha256(data).hexdigest()
        meta: dict[str, Any] = {
            "original_filename": original_filename,
            "file_size": len(data),
            "ocr_region_drafts": region_drafts,
            "ocr_base_text": base_text,
            "ocr_provider": ocr_engine.provider_name if ocr_engine.enabled else "none",
            "image_count": image_no,
        }
        return ExtractedDocument(
            text=full_text,
            content_hash=content_hash,
            source_uri=f"upload://{content_hash}",
            mime=DOCX_MIME,
            raw_bytes=data,
            raw_filename="raw.docx",
            metadata=meta,
        )
